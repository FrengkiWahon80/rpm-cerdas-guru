import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIG HALAMAN ---
st.set_page_config(page_title="RPM DEEP LEARNING AI", page_icon="🧠", layout="wide")

# --- LOGIKA AI: GENERATOR KONTEN DETAIL ---
def generate_deep_learning_content(topik, subtopik):
    return {
        "awal": (
            f"1. **Orientasi**: Guru menyapa peserta didik dengan hangat, berdoa bersama, dan mengecek kehadiran.\n"
            f"2. **Apersepsi**: Guru mengaitkan materi {topik} dengan pengalaman hidup sehari-hari siswa atau materi sebelumnya.\n"
            f"3. **Motivasi**: Guru memberikan gambaran manfaat mempelajari {subtopik} dalam kehidupan beriman dan bermasyarakat.\n"
            f"4. **Penyampaian Tujuan**: Guru menjelaskan kompetensi yang harus dicapai dan langkah pembelajaran mendalam yang akan dilalui."
        ),
        "inti": {
            "surface": (
                f"**Surface (Pemerolehan)**:\n"
                f"- Peserta didik membaca teks Kitab Suci/dokumen terkait {topik}.\n"
                f"- Mengidentifikasi fakta-fakta dasar, kosakata kunci, dan informasi tersurat dalam {subtopik}.\n"
                f"- Guru melakukan tanya jawab singkat untuk memastikan pemahaman dasar."
            ),
            "deep": (
                f"**Deep (Pengolahan)**:\n"
                f"- Peserta didik bekerja dalam kelompok untuk menganalisis makna mendalam dari {subtopik}.\n"
                f"- Diskusi kritis: Menghubungkan ajaran dengan tantangan zaman saat ini (Problem Based Learning).\n"
                f"- Guru memfasilitasi debat atau diskusi panel tentang nilai-nilai moral yang terkandung."
            ),
            "transfer": (
                f"**Transfer (Penerapan)**:\n"
                f"- Peserta didik membuat proyek kreatif (video pendek/poster/artikel) yang menunjukkan penerapan nilai {topik}.\n"
                f"- Menyusun rencana aksi nyata yang akan dilakukan di lingkungan sekolah atau rumah.\n"
                f"- Presentasi hasil karya untuk mendapatkan umpan balik dari rekan sejawat."
            )
        },
        "penutup": (
            f"1. **Refleksi**: Peserta didik merenungkan apa yang paling menyentuh hati selama pembelajaran {subtopik}.\n"
            f"2. **Umpan Balik**: Guru memberikan penguatan dan apresiasi atas proses belajar yang telah dilalui.\n"
            f"3. **Kesimpulan & Komitmen**: Guru dan siswa merangkum poin utama dan membuat komitmen bersama untuk menerapkannya.\n"
            f"4. **Doa Penutup**: Menutup kegiatan dengan doa syukur."
        ),
        "rubrik": (
            "1. **Pemahaman Materi (40%)**: Mampu menjelaskan konsep dan dasar biblis dengan tepat.\n"
            "2. **Analisis Kritis (30%)**: Mampu menghubungkan ajaran dengan realitas hidup secara mendalam.\n"
            "3. **Kualitas Produk (20%)**: Kreativitas dan kejelasan pesan dalam hasil karya (Transfer).\n"
            "4. **Sikap/Kolaborasi (10%)**: Keaktifan, etika, dan kerjasama dalam kelompok."
        )
    }

# --- FUNGSI EKSPOR WORD (SESUAI TEMPLATE PDF) ---
def set_cell_background(cell, fill_color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tc_pr.append(shd)

def export_word_rpm(data):
    doc = Document()
    
    # Judul Utama
    title = doc.add_paragraph()
    run = title.add_run("PERENCANAAN PEMBELAJARAN MENDALAM (DEEP LEARNING)")
    run.font.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tabel Identitas Atas
    tbl_top = doc.add_table(rows=5, cols=2)
    tbl_top.style = 'Table Grid'
    rows_top = [
        ("SEKOLAH", f": {data['sekolah']}"),
        ("NAMA GURU", f": {data['guru']}"),
        ("MATA PELAJARAN", f": {data['mapel']}"),
        ("KELAS / SEMESTER", f": {data['kelas']}"),
        ("ALOKASI WAKTU", f": {data['durasi']}")
    ]
    for i, (k, v) in enumerate(rows_top):
        tbl_top.rows[i].cells[0].text = k
        tbl_top.rows[i].cells[1].text = v

    doc.add_paragraph() # Spacer

    # Tabel Utama Identifikasi, Desain, Pengalaman
    tbl_main = doc.add_table(rows=15, cols=3)
    tbl_main.style = 'Table Grid'
    
    def fill_row(idx, col1, col2, col3):
        tbl_main.rows[idx].cells[0].text = col1
        tbl_main.rows[idx].cells[1].text = col2
        tbl_main.rows[idx].cells[2].text = col3
        set_cell_background(tbl_main.rows[idx].cells[1], "F2F2F2")

    # Bagian IDENTIFIKASI
    fill_row(0, "IDENTIFIKASI", "Peserta Didik", data['karakteristik'])
    fill_row(1, "", "Materi Pelajaran", f"Topik Utama: {data['topik']}\nSub-topik: {data['subtopik']}")
    fill_row(2, "", "Dimensi Profil Lulusan", data['profil'])
    fill_row(3, "", "Capaian Pembelajaran", data['cp'])
    fill_row(4, "", "Lintas Disiplin Ilmu", data['lintas'])
    fill_row(5, "", "Tujuan Pembelajaran", data['tujuan'])

    # Bagian DESAIN PEMBELAJARAN
    fill_row(6, "DESAIN PEMBELAJARAN", "Praktik Pedagogis", data['pedagogis'])
    fill_row(7, "", "Kemitraan Pembelajaran", data['kemitraan'])
    fill_row(8, "", "Lingkungan Belajar", data['lingkungan'])
    fill_row(9, "", "Pemanfaatan Digital", data['digital'])

    # Bagian PENGALAMAN BELAJAR
    fill_row(10, "PENGALAMAN BELAJAR", "Awal (15-20 menit)", data['p_awal'])
    fill_row(11, "", "Inti (60-80 menit)", f"{data['p_inti_surf']}\n\n{data['p_inti_deep']}\n\n{data['p_inti_tran']}")
    fill_row(12, "", "Penutup (15-20 menit)", data['p_penutup'])

    # Bagian ASESMEN
    fill_row(13, "ASESMEN", "Teknik & Instrumen", data['as_teknik'])
    fill_row(14, "", "Rubrik Penilaian", data['as_rubrik'])

    # Tanda Tangan
    doc.add_paragraph("\n")
    ttd = doc.add_table(rows=3, cols=2)
    ttd.rows[0].cells[0].text = "Mengetahui,\nKepala Sekolah"
    ttd.rows[0].cells[1].text = f"Merdeka, ................ 2025\nGuru Mata Pelajaran"
    ttd.rows[2].cells[0].text = f"{data['kepsek']}\nNIP. {data['nip_kepsek']}"
    ttd.rows[2].cells[1].text = f"{data['guru']}\nNIP. {data['nip_guru']}"

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- INTERFACE UTAMA STREAMLIT ---
st.title("📄 RPM DEEP LEARNING GENERATOR")
st.markdown("Sesuai Template SMPN Tujuh Maret Hadakewa")

with st.sidebar:
    st.header("📋 Administrasi")
    sekolah = st.text_input("Sekolah", "SMPN Tujuh Maret Hadakewa")
    guru = st.text_input("Nama Guru", "Daniel Florensius Lako Wahon, S.S")
    nip_guru = st.text_input("NIP Guru", "19801032024211002")
    kepsek = st.text_input("Kepala Sekolah", "Fransiskus Bernardus Kedang Kaona, S.Fl")
    nip_kepsek = st.text_input("NIP Kepala Sekolah", "19800132006041015")
    mapel = st.text_input("Mata Pelajaran", "Agama Katolik dan Budi Pekerti")
    kelas = st.text_input("Kelas / Semester", "IX / 1")
    durasi = st.text_input("Alokasi Waktu", "2 JP (2 x 40 Menit)")

st.subheader("🔍 1. Identifikasi & Desain")
c1, c2 = st.columns(2)
with c1:
    topik = st.text_input("Topik Utama", "Gereja yang Beriman")
    subtopik = st.text_input("Sub-topik", "Peran Serta dalam Hirarki")
    profil = st.multiselect("Dimensi Profil Pelajar Pancasila", 
                           ["Beriman & Bertakwa", "Mandiri", "Bernalar Kritis", "Kreatif", "Gotong Royong", "Kebinekaan Global"],
                           default=["Beriman & Bertakwa", "Bernalar Kritis"])
    lintas = st.text_input("Lintas Disiplin Ilmu", "PPKn (Struktur Organisasi), Bahasa Indonesia (Literasi Kitab Suci)")
with c2:
    pedagogis = st.selectbox("Praktik Pedagogis", ["Pembelajaran Berbasis Masalah", "Pembelajaran Berbasis Proyek", "Inkuiri", "Kontekstual"])
    digital = st.text_input("Pemanfaatan Digital", "Video Pembelajaran, Canva, Google Form untuk Refleksi")
    kemitraan = st.text_area("Kemitraan Pembelajaran", "Orang tua (diskusi iman), Tokoh Agama (Narasumber)", height=68)

st.divider()
st.subheader("🚀 2. Pengalaman Belajar (Detail)")

ai_content = generate_deep_learning_content(topik, subtopik)

tab_awal, tab_inti, tab_penutup = st.tabs(["Awal (15-20 Menit)", "Inti (60-80 Menit)", "Penutup (15-20 Menit)"])

with tab_awal:
    p_awal = st.text_area("Aktivitas Pembuka", ai_content['awal'], height=150)

with tab_inti:
    st.info("Strategi Deep Learning: Surface ➔ Deep ➔ Transfer")
    col_s, col_d, col_t = st.columns(3)
    p_surf = col_s.text_area("Surface (Pemerolehan)", ai_content['inti']['surface'], height=250)
    p_deep = col_d.text_area("Deep (Pengolahan)", ai_content['inti']['deep'], height=250)
    p_tran = col_t.text_area("Transfer (Penerapan)", ai_content['inti']['transfer'], height=250)

with tab_penutup:
    p_penutup = st.text_area("Aktivitas Penutup", ai_content['penutup'], height=150)

st.divider()
st.subheader("📊 3. Asesmen & Rubrik")
ca, cb = st.columns(2)
with ca:
    as_teknik = st.text_area("Teknik & Instrumen Penilaian", 
                            "1. Penilaian Diri (As Learning)\n2. Observasi Diskusi (For Learning)\n3. Portofolio/Produk Kreatif (Of Learning)")
with cb:
    as_rubrik = st.text_area("Rubrik Penilaian Detail", ai_content['rubrik'], height=150)

# Collect Data for Export
data_rpm = {
    "sekolah": sekolah, "guru": guru, "nip_guru": nip_guru, "kepsek": kepsek, "nip_kepsek": nip_kepsek,
    "mapel": mapel, "kelas": kelas, "durasi": durasi,
    "karakteristik": "Peserta didik memiliki latar belakang iman yang beragam namun memerlukan bimbingan dalam analisis kritis teks Kitab Suci.",
    "topik": topik, "subtopik": subtopik, "profil": ", ".join(profil),
    "cp": "Peserta didik memahami konteks ajaran Gereja dan mampu menerapkannya dalam kehidupan menggereja secara aktif.",
    "lintas": lintas, "tujuan": f"Peserta didik mampu menganalisis {subtopik} dan menghasilkan rencana aksi nyata.",
    "pedagogis": pedagogis, "kemitraan": kemitraan, "lingkungan": "Ruang kelas yang nyaman dan Ruang Digital (G-Classroom)",
    "digital": digital, "p_awal": p_awal, "p_inti_surf": p_surf, "p_inti_deep": p_deep, "p_inti_tran": p_tran,
    "p_penutup": p_penutup, "as_teknik": as_teknik, "as_rubrik": as_rubrik
}

if st.button("📝 Generate & Download RPM Mendalam (.docx)", type="primary"):
    file_word = export_word_rpm(data_rpm)
    st.download_button(
        label="📥 Unduh Sekarang",
        data=file_word,
        file_name=f"RPM_Mendalam_{topik}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
